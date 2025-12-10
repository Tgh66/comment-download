import streamlit as st
import yt_dlp
import os
import pandas as pd
import io
from datetime import datetime

# 引入导出所需的库
from docx import Document
from fpdf import FPDF

# --- 页面设置 ---
st.set_page_config(page_title="通用视频下载与数据分析工具", layout="wide")

st.title("📺 全能视频助手：下载 + 数据分析")
st.markdown("支持 Bilibili / YouTube / 抖音 | 自动抓取点赞数并排序")

# --- 初始化 Session State (用于存储抓取到的视频信息) ---
if 'video_data' not in st.session_state:
    st.session_state.video_data = []

# --- 核心工具函数 ---

def get_video_info_and_download(url, save_path="downloads"):
    """
    核心逻辑：既下载视频，又提取元数据用于后续排序分析
    """
    if not os.path.exists(save_path):
        os.makedirs(save_path)
    
    ydl_opts = {
        'outtmpl': f'{save_path}/%(title)s.%(ext)s',
        'format': 'best',
        # 安静模式，只打印关键信息
        'quiet': True,
        'no_warnings': True,
    }

    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            # 1. 提取信息 (不下载，先获取元数据)
            info_dict = ydl.extract_info(url, download=False)
            
            # 获取关键字段，如果没有则默认为0或未知
            video_title = info_dict.get('title', 'Unknown Title')
            like_count = info_dict.get('like_count', 0) # 核心：获取点赞数
            view_count = info_dict.get('view_count', 0)
            uploader = info_dict.get('uploader', 'Unknown')
            webpage_url = info_dict.get('webpage_url', url)
            
            # 处理点赞数可能为None的情况
            if like_count is None: like_count = 0

            # 2. 存入 Session State 用于排序和导出
            # 检查是否重复添加
            if not any(d['url'] == webpage_url for d in st.session_state.video_data):
                st.session_state.video_data.append({
                    "title": video_title,
                    "likes": int(like_count),
                    "views": int(view_count),
                    "uploader": uploader,
                    "url": webpage_url,
                    "scraped_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                })
            
            # 3. 执行下载 (保持原有功能)
            ydl.download([url])
            
            return True, f"下载成功: {video_title}", info_dict
            
    except Exception as e:
        return False, f"发生错误: {str(e)}", None

# --- 导出函数 ---

def to_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')
    return output.getvalue()

def to_word(df):
    doc = Document()
    doc.add_heading('视频数据分析报告', 0)
    
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)
    
    # 内容
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

def to_pdf(df):
    # 注意：标准FPDF不支持中文字符，需注册字体。
    # 这里为了代码通用性，使用英文表头或简单处理。
    # 实际生产环境建议下载支持中文的 .ttf 文件并加载。
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    
    pdf.cell(200, 10, txt="Video Data Report", ln=1, align='C')
    
    # 简单列表展示
    for index, row in df.iterrows():
        line = f"Title: {row['title'][:30]}... | Likes: {row['likes']} | Uploader: {row['uploader']}"
        # 移除非ASCII字符以防报错 (FPDF 默认限制)
        line = line.encode('latin-1', 'replace').decode('latin-1')
        pdf.cell(0, 10, txt=line, ln=1)
        
    return pdf.output(dest='S').encode('latin-1')

# --- 界面布局 ---

# 1. 输入与下载区 (保持原有逻辑)
st.subheader("1. 视频下载与抓取")
url_input = st.text_input("请输入视频链接 (支持 Bilibili, YouTube, Douyin):")
if st.button("开始下载并抓取数据"):
    if url_input:
        with st.spinner("正在解析并下载中..."):
            success, msg, info = get_video_info_and_download(url_input)
            if success:
                st.success(msg)
                st.balloons()
            else:
                st.error(msg)
    else:
        st.warning("请先输入链接")

st.divider()

# 2. 数据分析与导出区 (新增功能)
if st.session_state.video_data:
    st.subheader("2. 数据分析与导出")
    
    # 创建 DataFrame
    df = pd.DataFrame(st.session_state.video_data)
    
    # --- 排序控制 ---
    col1, col2 = st.columns([1, 3])
    with col1:
        sort_order = st.radio(
            "按照点赞数排序:",
            ('默认 (按时间)', '点赞数 (从高到低 ⬇️)', '点赞数 (从低到高 ⬆️)')
        )
    
    # 处理排序逻辑
    if sort_order == '点赞数 (从高到低 ⬇️)':
        df = df.sort_values(by='likes', ascending=False)
    elif sort_order == '点赞数 (从低到高 ⬆️)':
        df = df.sort_values(by='likes', ascending=True)
    
    # 展示表格
    st.dataframe(
        df, 
        column_config={
            "url": st.column_config.LinkColumn("视频链接"),
            "likes": st.column_config.NumberColumn("点赞数", format="%d 👍"),
            "views": st.column_config.NumberColumn("播放量", format="%d 👁️")
        },
        use_container_width=True
    )
    
    st.write("---")
    st.write("### 📥 导出数据")
    
    # --- 导出按钮 ---
    d_col1, d_col2, d_col3, d_col4 = st.columns(4)
    
    # 1. CSV
    csv_data = df.to_csv(index=False).encode('utf-8-sig')
    d_col1.download_button("下载 CSV", data=csv_data, file_name="video_data.csv", mime="text/csv")
    
    # 2. Excel
    excel_data = to_excel(df)
    d_col2.download_button("下载 Excel", data=excel_data, file_name="video_data.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    
    # 3. Word
    word_data = to_word(df)
    d_col3.download_button("下载 Word", data=word_data, file_name="video_data.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    # 4. PDF
    pdf_data = to_pdf(df)
    d_col4.download_button("下载 PDF", data=pdf_data, file_name="video_data.pdf", mime="application/pdf")

else:
    st.info("暂无数据。请在上方输入链接并下载，数据将自动添加至此处。")
